#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据动物名录.xlsx生成与动物名录.docx格式相同的Word文档（横向A4）
"""
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import os
import argparse
from pathlib import Path


def set_chinese_font(run, font_name='宋体', size=10.5, bold=False, italic=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_landscape(doc):
    """设置页面为横向A4"""
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

def add_heading_paragraph(doc, text, font_name='黑体', size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加标题段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, font_name, size, bold)
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    return p

def add_note_paragraph(doc, text, font_name='宋体', size=10.5):
    """添加注脚段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, font_name, size)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def setup_table_columns(table, widths):
    """设置表格列宽"""
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

def merge_row_cells(row, start_col=0, end_col=None):
    """合并一行中从start_col到end_col的所有单元格"""
    if end_col is None:
        end_col = len(row.cells) - 1
    if start_col < end_col:
        row.cells[start_col].merge(row.cells[end_col])
    return row.cells[start_col]

def add_group_row(table, text, cols_count, bold=True, font_size=10.5):
    """添加分组行（目/科），合并所有列"""
    row = table.add_row()
    cell = merge_row_cells(row, 0, cols_count - 1)
    run = cell.paragraphs[0].add_run(text)
    set_chinese_font(run, '黑体', font_size, bold=bold)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row.height = Pt(20)
    return row

def add_species_row_amphibian_reptile(table, idx, chinese_name, latin_name, habitat, zone, protection, san_you, source) -> Any:
    """添加两栖/爬行类物种行"""
    row = table.add_row()
    cell0 = row.cells[0]
    p0 = cell0.paragraphs[0]
    run_text = f"{idx}.  {chinese_name}  "
    run1 = p0.add_run(run_text)
    set_chinese_font(run1, '宋体', 10.5)
    if pd.notna(latin_name) and str(latin_name).strip():
        run2 = p0.add_run(str(latin_name))
        set_chinese_font(run2, 'Times New Roman', 10.5, italic=True)

    row.cells[1].text = str(habitat) if pd.notna(habitat) else ""
    row.cells[2].text = str(zone) if pd.notna(zone) else ""
    row.cells[3].text = ""
    row.cells[4].text = str(protection) if pd.notna(protection) else "—"
    row.cells[5].text = "√" if pd.notna(san_you) and str(san_you).strip() in ['√', '1', 'True'] else "—"
    row.cells[6].text = ""
    row.cells[7].text = str(source) if pd.notna(source) else "—"

    for i in range(2, 8):
        cell = row.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '宋体', 10.5)
    row.height = Pt(28)
    return row

def add_species_row_bird(table, chinese_name, latin_name, residence, zone, habitat, protection, san_you, cites,source):
    """添加鸟类物种行"""
    row = table.add_row()
    cell0 = row.cells[0]
    p0 = cell0.paragraphs[0]
    run1 = p0.add_run(str(chinese_name) if pd.notna(chinese_name) else "")
    set_chinese_font(run1, '宋体', 10.5)
    if pd.notna(latin_name) and str(latin_name).strip():
        run2 = p0.add_run(f"\n{str(latin_name)}")
        set_chinese_font(run2, 'Times New Roman', 10.5, italic=True)

    row.cells[1].text = str(residence) if pd.notna(residence) else ""
    row.cells[2].text = str(zone) if pd.notna(zone) else ""
    row.cells[3].text = str(habitat) if pd.notna(habitat) else ""
    row.cells[4].text = str(protection) if pd.notna(protection) else "—"
    row.cells[5].text = ""
    row.cells[6].text = ""
    row.cells[7].text = "√" if pd.notna(san_you) and str(san_you).strip() in ['√', '1', 'True'] else "—"
    row.cells[8].text = str(cites) if pd.notna(cites) else ""
    row.cells[9].text = str(source) if pd.notna(source) else ""

    for i in range(1, 3):
        cell = row.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for i in range(4, 10):
        cell = row.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '宋体', 10.5)
    row.height = Pt(28)
    return row

def add_species_row_mammal(table, idx, chinese_name, latin_name, zone, protection, san_you, source):
    """添加兽类物种行"""
    row = table.add_row()
    cell0 = row.cells[0]
    p0 = cell0.paragraphs[0]
    run1 = p0.add_run(str(chinese_name) if pd.notna(chinese_name) else "")
    set_chinese_font(run1, '宋体', 10.5)
    if pd.notna(latin_name) and str(latin_name).strip():
        run2 = p0.add_run(f"\n{str(latin_name)}")
        set_chinese_font(run2, 'Times New Roman', 10.5, italic=True)

    row.cells[1].text = str(zone) if pd.notna(zone) else ""
    row.cells[2].text = ""
    row.cells[3].text = str(protection) if pd.notna(protection) else "—"
    row.cells[4].text = "√" if pd.notna(san_you) and str(san_you).strip() in ['√', '1', 'True'] else "—"
    row.cells[5].text = ""
    row.cells[6].text = str(source) if pd.notna(source) else ""


    for i in range(1, 7):
        cell = row.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, '宋体', 10.5)
    row.height = Pt(28)
    return row

def create_amphibian_table(doc, df_group):
    """创建两栖类表格（附录4-1）"""
    cols = ['目、科、种名', '生境', '区系\n类型', '数量\n等级', '保护\n等级', '三有\n动物', '红色\n名录', '来源']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(col_name)
        set_chinese_font(run, '宋体', 10.5, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 横向页面列宽适当加宽
    widths = [Cm(5.5), Cm(9.0), Cm(2.2), Cm(1.5), Cm(2.4), Cm(1.5), Cm(1.5), Cm(1.5)]
    setup_table_columns(table, widths)

    orders = df_group.groupby('目')
    order_idx = 0
    for order_name, order_df in orders:
        order_idx += 1
        order_latin = order_df['Order'].iloc[0] if pd.notna(order_df['Order'].iloc[0]) else ""
        add_group_row(table, f"{order_idx}.  {order_name}{order_latin.upper()}", len(cols), bold=True)

        families = order_df.groupby('科')
        family_idx = 0
        for family_name, family_df in families:
            family_idx += 1
            family_latin = family_df['Family'].iloc[0] if pd.notna(family_df['Family'].iloc[0]) else ""
            add_group_row(table, f"({family_idx})  {family_name} {family_latin}", len(cols), bold=True)

            species_idx = 0
            for _, row in family_df.iterrows():
                species_idx += 1
                add_species_row_amphibian_reptile(
                    table, species_idx,
                    row['中文名'], row['拉丁名'],
                    row['生境'], row['区系类型'],
                    row['保护级别'], row['三有动物'],
                    row['来源']
                )
    return table

def create_reptile_table(doc, df_group):
    """创建爬行类表格（附录4-2）"""
    cols = ['目、科、种名', '生境', '区系\n类型', '数量\n等级', '保护\n等级', '三有\n动物', '红色\n名录', '来源']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(col_name)
        set_chinese_font(run, '宋体', 10.5, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    widths = [Cm(5.5), Cm(8.0), Cm(2.2), Cm(1.8), Cm(2.2), Cm(1.8), Cm(1.8), Cm(1.8)]
    setup_table_columns(table, widths)

    orders = df_group.groupby('目')
    order_idx = 0
    for order_name, order_df in orders:
        order_idx += 1
        order_latin = order_df['Order'].iloc[0] if pd.notna(order_df['Order'].iloc[0]) else ""
        add_group_row(table, f"{order_idx}.  {order_name}{order_latin.upper()}", len(cols), bold=True)

        families = order_df.groupby('科')
        family_idx = 0
        for family_name, family_df in families:
            family_idx += 1
            family_latin = family_df['Family'].iloc[0] if pd.notna(family_df['Family'].iloc[0]) else ""
            add_group_row(table, f"({family_idx})  {family_name} {family_latin}", len(cols), bold=True)

            species_idx = 0
            for _, row in family_df.iterrows():
                species_idx += 1
                add_species_row_amphibian_reptile(
                    table, species_idx,
                    row['中文名'], row['拉丁名'],
                    row['生境'], row['区系类型'],
                    row['保护级别'], row['三有动物'],
                    row['来源']
                )
    return table

def create_bird_table(doc, df_group):
    """创建鸟类表格（附录4-3）"""
    cols = ['中文名、拉丁名', '居留型', '区系', '生境', '保护等级', '数量等级', '红色名录', '三有动物', 'CITES', '来源']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(col_name)
        set_chinese_font(run, '宋体', 10.5, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    widths = [Cm(5.5), Cm(1.8), Cm(1.8), Cm(7.0), Cm(2.2), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8)]
    setup_table_columns(table, widths)

    orders = df_group.groupby('目')
    order_idx = 0
    for order_name, order_df in orders:
        order_idx += 1
        order_latin = order_df['Order'].iloc[0] if pd.notna(order_df['Order'].iloc[0]) else ""
        add_group_row(table, f"{order_idx}. {order_name}{order_latin.upper()}", len(cols), bold=True)

        families = order_df.groupby('科')
        family_idx = 0
        for family_name, family_df in families:
            family_idx += 1
            family_latin = family_df['Family'].iloc[0] if pd.notna(family_df['Family'].iloc[0]) else ""
            add_group_row(table, f"({family_idx})  {family_name} {family_latin}", len(cols), bold=False)

            species_idx = 0
            for _, row in family_df.iterrows():
                species_idx += 1
                add_species_row_bird(
                    table,
                    row['中文名'], row['拉丁名'],
                    row['居留型'], row['区系类型'],
                    row['生境'], row['保护级别'],
                    row['三有动物'], row['濒危野生动植物种国际贸易公约'],
                    row['来源']
                )
    return table

def create_mammal_table(doc, df_group):
    """创建兽类表格（附录4-4）"""
    cols = ['种中文名', '拉丁种名', '区系', '数量等级', '保护等级', '三有动物', '红色名录', '来源']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(col_name)
        set_chinese_font(run, '宋体', 10.5, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    widths = [Cm(4.5), Cm(5.5), Cm(2.2), Cm(1.8), Cm(2.2), Cm(1.8), Cm(1.8), Cm(1.8)]
    setup_table_columns(table, widths)

    orders = df_group.groupby('目')
    order_idx = 0
    for order_name, order_df in orders:
        order_idx += 1
        order_latin = order_df['Order'].iloc[0] if pd.notna(order_df['Order'].iloc[0]) else ""
        add_group_row(table, f"{order_idx}.  {order_name}{order_latin.upper()}", len(cols), bold=True)

        families = order_df.groupby('科')
        family_idx = 0
        for family_name, family_df in families:
            family_idx += 1
            family_latin = family_df['Family'].iloc[0] if pd.notna(family_df['Family'].iloc[0]) else ""
            add_group_row(table, f"({family_idx})  {family_name} {family_latin}", len(cols), bold=True)

            species_idx = 0
            for _, row in family_df.iterrows():
                species_idx += 1
                add_species_row_mammal(
                    table, species_idx,
                    row['中文名'], row['拉丁名'],
                    row['区系类型'], row['保护级别'],
                    row['三有动物'], row['来源']
                )
    return table

def main():
    parser = argparse.ArgumentParser(description='利用动物名录的excel表格生成word版')
    # 输入输出文件参数
    parser.add_argument('--work_dir', default='D:/EcoAgentProject',
                        help='工作文件路径')
    parser.add_argument('--input_file', default='动物名录.xlsx',
                        help='动物名录的excel表格')
    parser.add_argument('--output_file', default='动物名录.docx',
                        help='动物名录word版')

    args = parser.parse_args()

    work_dir = Path(args.work_dir).resolve()  # resolve() 会统一格式并转为绝对路径
    excel_path = work_dir / args.input_file
    output_path = work_dir / args.output_file

    df = pd.read_excel(excel_path)
    doc = Document()

    # 设置横向A4
    set_landscape(doc)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # 附录4-1 两栖类
    add_heading_paragraph(doc, '附录4-1 两栖类名录', font_name='黑体', size=14, bold=True)
    df_amphibia = df[df['纲'] == '两栖纲'].copy()
    if not df_amphibia.empty:
        create_amphibian_table(doc, df_amphibia)
        add_note_paragraph(doc,
            "注：分类系统参照《中国两栖类信息系统》（中国科学院昆明动物研究所，2024）；\n"
            "三有动物：《国家保护的有益的或者有重要经济、科学研究价值的陆生野生动物名录》；\n"
            "红色名录栏： VU---易危、NT---近危、LC---无危；")
    doc.add_page_break()

    # 附录4-2 爬行类
    add_heading_paragraph(doc, '附录4-2 爬行类名录', font_name='黑体', size=14, bold=True)
    df_reptile = df[df['纲'] == '爬行纲'].copy()
    if not df_reptile.empty:
        create_reptile_table(doc, df_reptile)
        add_note_paragraph(doc,
            "注：分类系统参照《中国爬行纲动物分类厘定》（蔡波等，2015）；\n"
            "红色名录栏：EN---濒危、VU---易危、LC---无危；\n"
            "三有动物：《国家保护的有益的或者有重要经济、科学研究价值的陆生野生动物名录》")
    doc.add_page_break()

    # 附录4-3 鸟类
    add_heading_paragraph(doc, '附录4-3 鸟类名录', font_name='黑体', size=14, bold=True)
    df_bird = df[df['纲'] == '鸟纲'].copy()
    if not df_bird.empty:
        create_bird_table(doc, df_bird)
        add_note_paragraph(doc,
            "注：①分类系统参照《中国鸟类分类与分布名录》（郑光美，2022）；\n"
            "②红色名录：EN---濒危、VU---易危、NT---近危、LC---无危；\n"
            "③三有动物：《国家保护的有益的或者有重要经济、科学研究价值的陆生野生动物名录》；\n"
            "④\"CITES\"栏：《濒危野生动植物种国际贸易公约》Ⅰ---附录Ⅰ，Ⅱ---附录Ⅱ；")
    doc.add_page_break()

    # 附录4-4 兽类
    add_heading_paragraph(doc, '附录4-4 兽类名录', font_name='黑体', size=14, bold=True)
    df_mammal = df[df['纲'] == '哺乳纲'].copy()
    if not df_mammal.empty:
        create_mammal_table(doc, df_mammal)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"横向文档已成功生成：{output_path}")

if __name__ == '__main__':
    main()